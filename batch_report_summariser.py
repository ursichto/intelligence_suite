import os
import html
import re
from openai import OpenAI
from PyPDF2 import PdfReader
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# === CONFIGURATION ===
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
INPUT_FOLDER = "input_reports"
OUTPUT_FOLDER = "summaries"
MODEL = "gpt-4o-mini"

USD_PER_1K_PROMPT = 0.00015
USD_PER_1K_COMPLETION = 0.00060
CHF_CONVERSION_RATE = 0.80

LOGO_PATH = "Transformate Logo.png"


# === HELPERS ===
def extract_text_from_pdf(path):
    """Extract text from a PDF."""
    reader = PdfReader(path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text.strip()


def summarise_text(text, filename):
    """Generate executive summary and return summary + token stats."""
    prompt = f"""
    You are an expert financial analyst.
    Summarise the following document into a concise executive summary (max 400 words).
    Include:
    - Main purpose or regulatory intent
    - Key operational or compliance impacts
    - 3–5 recommended next steps for an executive team.

    Document: {text[:12000]}
    """

    print(f"→ Summarising {filename}...")
    response = client.chat.completions.create(
        model=MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4,
    )

    summary = response.choices[0].message.content.strip()
    usage = response.usage
    prompt_tokens = usage.prompt_tokens
    completion_tokens = usage.completion_tokens
    total_tokens = usage.total_tokens

    usd_cost = (
        (prompt_tokens / 1000) * USD_PER_1K_PROMPT
        + (completion_tokens / 1000) * USD_PER_1K_COMPLETION
    )
    chf_cost = usd_cost * CHF_CONVERSION_RATE

    return summary, prompt_tokens, completion_tokens, total_tokens, usd_cost, chf_cost


# === DOCX CREATION ===
def save_to_docx(
    summary_text,
    filename,
    source_document=None,
    prompt_tokens=None,
    completion_tokens=None,
    total_tokens=None,
    usd_cost=None,
    chf_cost=None,
):
    """Save summary text to a Word document with cover page and accurate logo scaling (no blank page)."""
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime
    from zoneinfo import ZoneInfo
    from PIL import Image as PILImage

    src_name = os.path.basename(source_document) if source_document else ""
    first_line = next((l.strip() for l in summary_text.split("\n") if l.strip()), "")
    clean_first = first_line.replace("**", "").strip("* ").strip()
    m = re.search(r"Executive Summary\s*:\s*(.*)", clean_first, flags=re.I)
    exec_title = m.group(1).strip() if m and m.group(1).strip() else clean_first
    if not exec_title:
        exec_title = os.path.splitext(src_name)[0]

    dt = datetime.now(ZoneInfo("Europe/Zurich")).strftime("%B %d %Y, %H:%M CET")

    doc = Document()

    # --- COVER PAGE ---
    if os.path.exists(LOGO_PATH):
        with PILImage.open(LOGO_PATH) as img:
            width, height = img.size
            target_width = Inches(2.2)
            aspect_ratio = height / width
            target_height = target_width * aspect_ratio
        doc.add_picture(LOGO_PATH, width=target_width, height=target_height)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

    p_title = doc.add_paragraph("Executive Summary")
    p_title.style = doc.styles["Title"]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("")
    doc.add_paragraph("Executive Summary Title:").runs[0].bold = True
    doc.add_paragraph(exec_title)
    doc.add_paragraph("")
    doc.add_paragraph(f"Generated on: {dt}")
    doc.add_paragraph(f"Source Document: {src_name}")
    doc.add_paragraph("Prepared by: Transformate Consulting | AI Generated Summary")

    # --- AI Processing Summary on cover page ---
    if all(v is not None for v in [prompt_tokens, completion_tokens, total_tokens, usd_cost, chf_cost]):
        doc.add_paragraph("")
        p = doc.add_paragraph("AI Processing Summary:")
        p.runs[0].bold = True
        doc.add_paragraph(
            f"Prompt Tokens: {prompt_tokens:,} | Completion Tokens: {completion_tokens:,} | Total: {total_tokens:,}"
        )
        doc.add_paragraph(
            f"Estimated Cost: USD ${usd_cost:.4f} ≈ CHF {chf_cost:.4f}"
        )

    doc.add_paragraph("")
    doc.add_paragraph("Contact:")
    doc.add_paragraph("Tony Ursich")
    doc.add_paragraph("Chief Information Officer | Transformate Consulting")
    doc.add_paragraph("E: tony.ursich@transformate.ch")
    doc.add_paragraph("T: +41 76 577 1165")
    doc.add_paragraph("W: www.transformate.ch")

    # ✅ Instead of a forced page break (which causes a blank page),
    # insert a few line spaces and then start the summary heading.
    doc.add_paragraph("\n" * 2)

    # --- SUMMARY PAGES ---
    doc.add_heading("Executive Summary\n", level=1)

    # Remove redundant "Executive Summary:" line from body
    summary_text = re.sub(r"^(\*\*)?\s*Executive Summary\s*:?.*\n?", "", summary_text, flags=re.I)

    paragraphs = [p.strip() for p in summary_text.split("\n") if p.strip()]
    for para in paragraphs:
        p = doc.add_paragraph()
        parts = re.split(r"(\*\*.+?\*\*)", para)
        for part in parts:
            clean_text = part.replace("**", "")
            run = p.add_run(clean_text)
            if part.startswith("**") and part.endswith("**"):
                run.bold = True
            run.font.size = Pt(11)

    doc.save(filename)

# === PDF CREATION ===
def save_to_pdf(
    summary_text,
    filename,
    source_document=None,
    title="Executive Summary",
    prompt_tokens=None,
    completion_tokens=None,
    total_tokens=None,
    usd_cost=None,
    chf_cost=None,
):
    """Save summary text to a PDF with a left-aligned logo and Source Document metadata."""
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from datetime import datetime
    from zoneinfo import ZoneInfo
    from PIL import Image as PILImage
    import html, re

    src_name = os.path.basename(source_document) if source_document else ""
    first_line = next((l.strip() for l in summary_text.split("\n") if l.strip()), "")
    clean_first = first_line.replace("**", "").strip("* ").strip()
    m = re.search(r"Executive Summary\s*:\s*(.*)", clean_first, flags=re.I)
    exec_title = m.group(1).strip() if m and m.group(1).strip() else clean_first
    if not exec_title:
        exec_title = os.path.splitext(src_name)[0]
    dt = datetime.now(ZoneInfo("Europe/Zurich")).strftime("%B %d %Y, %H:%M CET")

    # remove redundant "Executive Summary:" line from summary_text
    summary_text = re.sub(r"^(\*\*)?\s*Executive Summary\s*:?.*\n?", "", summary_text, flags=re.I)

    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        leftMargin=25 * mm,
        rightMargin=25 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()
    cover_title = ParagraphStyle("CoverTitle", parent=styles["Title"], fontName="Helvetica-Bold",
                                 fontSize=24, leading=28, spaceAfter=10, alignment=0)
    cover_label = ParagraphStyle("CoverLabel", parent=styles["BodyText"], fontName="Helvetica-Bold",
                                 fontSize=11, leading=14, spaceAfter=2, alignment=0)
    cover_text = ParagraphStyle("CoverText", parent=styles["BodyText"], fontName="Helvetica",
                                fontSize=11, leading=14, spaceAfter=8, alignment=0)
    heading = ParagraphStyle("Heading", parent=styles["Heading1"], fontName="Helvetica-Bold",
                             fontSize=14, leading=18, spaceAfter=10, spaceBefore=6)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica",
                          fontSize=11, leading=15, spaceAfter=10)
    list_item = ParagraphStyle("ListItem", parent=styles["BodyText"], fontName="Helvetica",
                               fontSize=11, leading=15, spaceAfter=10)

    def _md_to_html(text):
        t = html.escape(text)
        t = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", t)
        t = t.replace("\r\n", "\n").replace("\n", "<br/>")
        return t

    story = []
    if os.path.exists(LOGO_PATH):
        with PILImage.open(LOGO_PATH) as img:
            w, h = img.size
            target_width = 50 * mm
            target_height = target_width * (h / w)
        story.append(Image(LOGO_PATH, width=target_width, height=target_height, hAlign='LEFT'))
        story.append(Spacer(1, 10))

    story.append(Paragraph("Executive Summary", cover_title))
    story.append(HRFlowable(width="100%", thickness=0.6, color=colors.lightgrey))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Executive Summary Title:", cover_label))
    story.append(Paragraph(html.escape(exec_title), cover_text))
    story.append(Spacer(1, 6))
    story.append(Paragraph(f"Generated on: {html.escape(dt)}", cover_text))
    story.append(Paragraph(f"Source Document: {html.escape(src_name)}", cover_text))
    story.append(Paragraph("Prepared by: Transformate Consulting | AI Generated Summary", cover_text))

    # --- NEW: AI Processing Summary on cover page (if values provided) ---
    if all(v is not None for v in [prompt_tokens, completion_tokens, total_tokens, usd_cost, chf_cost]):
        story.append(Spacer(1, 8))
        story.append(Paragraph("AI Processing Summary:", cover_label))
        story.append(Paragraph(
            html.escape(f"Prompt Tokens: {prompt_tokens:,} | Completion Tokens: {completion_tokens:,} | Total: {total_tokens:,}"),
            cover_text
        ))
        story.append(Paragraph(
            html.escape(f"Estimated Cost: USD ${usd_cost:.4f} ≈ CHF {chf_cost:.4f}"),
            cover_text
        ))

    story.append(Spacer(1, 12))
    story.append(HRFlowable(width="100%", thickness=0.6, color=colors.lightgrey))
    story.append(Spacer(1, 8))
    story.append(Paragraph("Contact:", cover_label))
    story.append(Paragraph("Tony Ursich", cover_text))
    story.append(Paragraph("Chief Information Officer | Transformate Consulting", cover_text))
    story.append(Paragraph("E: tony.ursich@transformate.ch", cover_text))
    story.append(Paragraph("T: +41 76 577 1165", cover_text))
    story.append(Paragraph("W: www.transformate.ch", cover_text))
    story.append(PageBreak())

    # --- SUMMARY ---
    story.append(Paragraph(title, heading))
    for line in [l.strip() for l in summary_text.split("\n") if l.strip()]:
        clean_line = line.strip()
        if clean_line.startswith("**") and clean_line.endswith("**"):
            clean_line = clean_line.replace("**", "")
            story.append(Paragraph(f"<b>{html.escape(clean_line)}</b>", heading))
            story.append(Spacer(1, 6))
            continue
        if re.match(r"^[A-Z].+?:$", clean_line) or (
            "Purpose and Regulatory Intent" in clean_line
            or "Key Operational" in clean_line
            or "Recommended Next Steps" in clean_line
        ):
            story.append(Paragraph(f"<b>{html.escape(clean_line)}</b>", heading))
            story.append(Spacer(1, 8))
            continue
        if re.match(r"^\d+\.", clean_line):
            html_line = _md_to_html(clean_line)
            html_line = re.sub(r"</b>:", r":</b>", html_line)
            story.append(Paragraph(html_line, list_item))
            story.append(Spacer(1, 6))
            continue
        story.append(Paragraph(_md_to_html(clean_line), body))
        story.append(Spacer(1, 8))
    doc.build(story)


# === MAIN ===
if __name__ == "__main__":
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    pdf_files = [f for f in os.listdir(INPUT_FOLDER) if f.lower().endswith(".pdf")]
    if not pdf_files:
        print(f"No PDFs found in '{INPUT_FOLDER}'. Please add files and rerun.")
        exit()

    print(f"Found {len(pdf_files)} PDF(s). Starting batch summarisation...\n")
    grand_total_tokens = grand_total_usd = grand_total_chf = 0

    for pdf in pdf_files:
        input_path = os.path.join(INPUT_FOLDER, pdf)
        base_name = os.path.splitext(pdf)[0]
        text = extract_text_from_pdf(input_path)
        summary, pt, ct, tt, usd, chf = summarise_text(text, pdf)
        output_docx = os.path.join(OUTPUT_FOLDER, f"Summary_{base_name}.docx")
        output_pdf = os.path.join(OUTPUT_FOLDER, f"Summary_{base_name}.pdf")

        save_to_docx(
            summary, output_docx, source_document=input_path,
            prompt_tokens=pt, completion_tokens=ct, total_tokens=tt,
            usd_cost=usd, chf_cost=chf
        )

        save_to_pdf(
            summary, output_pdf, source_document=input_path,
            prompt_tokens=pt, completion_tokens=ct, total_tokens=tt,
            usd_cost=usd, chf_cost=chf
        )

        print(f"✅ {base_name}: {tt} tokens | ${usd:.4f} USD | CHF {chf:.4f}")
        grand_total_tokens += tt
        grand_total_usd += usd
        grand_total_chf += chf

    print("\n=== SUMMARY REPORT ===")
    print(f"Total tokens used: {grand_total_tokens:,}")
    print(f"Total cost: ${grand_total_usd:.4f} USD ≈ CHF {grand_total_chf:.4f}")
    print("\n🎉 All summaries saved in the 'summaries' folder.")
