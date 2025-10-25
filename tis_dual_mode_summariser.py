#
# Transformate Intelligence Suite (Dual Mode + Chatbase Integration)
#

#
# Below is a drop-in FastAPI web service that keeps your exact PDF formatting logic and adds dual mode:
#
# POST /summarise_demo → runs all 3 fixed PDFs you listed (reads them from input_reports/).
# POST /summarise_upload → accepts up to 3 user PDFs and returns summaries.
# GET /files/{path:path} → serves generated PDFs (so Carrd can link/download).
# GET / → healthcheck.
#

import os
import re
import html
import shutil
import requests
from typing import List, Optional
from datetime import datetime
from zoneinfo import ZoneInfo

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware

from PyPDF2 import PdfReader
from openai import OpenAI

# ReportLab / DOCX deps
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from PIL import Image as PILImage

# === CONFIGURATION ===
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
CHATBASE_API_KEY = os.getenv("CHATBASE_API_KEY")  # <-- must be set in Render dashboard
CHATBASE_BOT_ID = "eJ86MiNhODp3671KJoMTN"

if not OPENAI_API_KEY:
    raise RuntimeError("OPENAI_API_KEY environment variable is not set.")

client = OpenAI(api_key=OPENAI_API_KEY)
MODEL = "gpt-4o-mini"

USD_PER_1K_PROMPT = 0.00015
USD_PER_1K_COMPLETION = 0.00060
CHF_CONVERSION_RATE = 0.80

BASE_DIR = os.path.abspath(os.path.dirname(__file__))
INPUT_FOLDER = os.path.join(BASE_DIR, "input_reports")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "summaries")
os.makedirs(INPUT_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

LOGO_PATH = os.path.join(BASE_DIR, "Transformate Logo.png")

DEMO_FILES = [
    "FINMA rs 2011 02 20200101 - Capital buffer and capital planning - banks.pdf",
    "FINMA rs 2023 01 20221207 - Operational risks and resilience - banks.pdf",
    "FINMA rs 2018 02 - Duty to report securities transactions.pdf",
]

# ---------- Helpers ----------

def extract_text_from_pdf(path: str) -> str:
    """Extract text from a PDF."""
    reader = PdfReader(path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text.strip()

def summarise_text(text: str, filename: str):
    """Generate executive summary and return summary + token/cost stats."""
    prompt = f"""
    You are an expert financial analyst.
    Summarise the following document into a concise executive summary (max 400 words).
    Include:
    - Main purpose or regulatory intent
    - Key operational or compliance impacts
    - 3–5 recommended next steps for an executive team.

    Document: {text[:12000]}
    """

    response = client.chat.completions.create(
        model=MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4,
    )

    summary = response.choices[0].message.content.strip()
    usage = response.usage
    prompt_tokens = usage.prompt_tokens
    completion_tokens = usage.completion_tokens
    total_tokens = usage.total_tokens

    usd_cost = (prompt_tokens / 1000) * USD_PER_1K_PROMPT + (completion_tokens / 1000) * USD_PER_1K_COMPLETION
    chf_cost = usd_cost * CHF_CONVERSION_RATE

    return summary, prompt_tokens, completion_tokens, total_tokens, usd_cost, chf_cost


# === Chatbase Integration ===
def push_to_chatbase(title: str, summary_text: str):
    """Send generated summary text to Chatbase knowledge base for live querying."""
    if not CHATBASE_API_KEY:
        print("⚠️ CHATBASE_API_KEY not configured — skipping Chatbase push.")
        return False
    url = "https://www.chatbase.co/api/knowledge"
    headers = {
        "Authorization": f"Bearer {CHATBASE_API_KEY}",
        "Content-Type": "application/json",
    }
    payload = {
        "chatbotId": CHATBASE_BOT_ID,
        "documents": [
            {
                "title": title,
                "content": summary_text[:18000]  # safeguard against oversized payloads
            }
        ],
    }
    try:
        r = requests.post(url, headers=headers, json=payload, timeout=20)
        r.raise_for_status()
        print(f"✅ Summary '{title}' uploaded to Chatbase.")
        return True
    except Exception as e:
        print(f"⚠️ Failed to upload summary '{title}' to Chatbase: {e}")
        return False

# ---------- DOCX and PDF GENERATION
# ---------- DOCX (kept intact; not used by demo endpoints unless you want it) ----------

def save_to_docx(
    summary_text: str,
    filename: str,
    source_document: Optional[str] = None,
    prompt_tokens: Optional[int] = None,
    completion_tokens: Optional[int] = None,
    total_tokens: Optional[int] = None,
    usd_cost: Optional[float] = None,
    chf_cost: Optional[float] = None,
):
    """Save summary text to a Word document with cover page and accurate logo scaling (no blank page)."""
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    src_name = os.path.basename(source_document) if source_document else ""
    first_line = next((l.strip() for l in summary_text.split("\n") if l.strip()), "")
    clean_first = first_line.replace("**", "").strip("* ").strip()
    m = re.search(r"Executive Summary\s*:\s*(.*)", clean_first, flags=re.I)
    exec_title = m.group(1).strip() if m and m.group(1).strip() else clean_first
    if not exec_title:
        exec_title = os.path.splitext(src_name)[0]

    dt = datetime.now(ZoneInfo("Europe/Zurich")).strftime("%B %d %Y, %H:%M CET")

    doc = Document()

    # Cover logo
    if os.path.exists(LOGO_PATH):
        with PILImage.open(LOGO_PATH) as img:
            width, height = img.size
            target_width = Inches(2.2)
            aspect_ratio = height / width
            target_height = target_width * aspect_ratio
        doc.add_picture(LOGO_PATH, width=target_width, height=target_height)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

    p_title = doc.add_paragraph("Executive Summary")
    p_title.style = doc.styles["Title"]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("")
    doc.add_paragraph("Executive Summary Title:").runs[0].bold = True
    doc.add_paragraph(exec_title)
    doc.add_paragraph("")
    doc.add_paragraph(f"Generated on: {dt}")
    doc.add_paragraph(f"Source Document: {src_name}")
    doc.add_paragraph("Prepared by: Transformate Consulting | AI Generated Summary")

    if all(v is not None for v in [prompt_tokens, completion_tokens, total_tokens, usd_cost, chf_cost]):
        doc.add_paragraph("")
        p = doc.add_paragraph("AI Processing Summary:")
        p.runs[0].bold = True
        doc.add_paragraph(
            f"Prompt Tokens: {prompt_tokens:,} | Completion Tokens: {completion_tokens:,} | Total: {total_tokens:,}"
        )
        doc.add_paragraph(f"Estimated Cost: USD ${usd_cost:.4f} ≈ CHF {chf_cost:.4f}")

    doc.add_paragraph("")
    doc.add_paragraph("Contact:")
    doc.add_paragraph("Tony Ursich")
    doc.add_paragraph("Chief Information Officer | Transformate Consulting")
    doc.add_paragraph("E: tony.ursich@transformate.ch")
    doc.add_paragraph("T: +41 76 577 1165")
    doc.add_paragraph("W: www.transformate.ch")

    # No forced page break (prevents blank page); small gap then heading.
    doc.add_paragraph("\n" * 2)

    # Summary pages
    doc.add_heading("Executive Summary\n", level=1)

    # Remove redundant "Executive Summary:" line from body
    summary_text = re.sub(r"^(\*\*)?\s*Executive Summary\s*:?.*\n?", "", summary_text, flags=re.I)

    paragraphs = [p.strip() for p in summary_text.split("\n") if p.strip()]
    for para in paragraphs:
        p = doc.add_paragraph()
        parts = re.split(r"(\*\*.+?\*\*)", para)
        for part in parts:
            clean_text = part.replace("**", "")
            run = p.add_run(clean_text)
            if part.startswith("**") and part.endswith("**"):
                run.bold = True
            run.font.size = Pt(11)

    doc.save(filename)


# ---------- PDF (exact formatting preserved) ----------

def save_to_pdf(
    summary_text: str,
    filename: str,
    source_document: Optional[str] = None,
    title: str = "Executive Summary",
    prompt_tokens: Optional[int] = None,
    completion_tokens: Optional[int] = None,
    total_tokens: Optional[int] = None,
    usd_cost: Optional[float] = None,
    chf_cost: Optional[float] = None,
):
    """Save summary text to a PDF with a left-aligned logo and Source Document metadata."""
    src_name = os.path.basename(source_document) if source_document else ""
    first_line = next((l.strip() for l in summary_text.split("\n") if l.strip()), "")
    clean_first = first_line.replace("**", "").strip("* ").strip()
    m = re.search(r"Executive Summary\s*:\s*(.*)", clean_first, flags=re.I)
    exec_title = m.group(1).strip() if m and m.group(1).strip() else clean_first
    if not exec_title:
        exec_title = os.path.splitext(src_name)[0]
    dt = datetime.now(ZoneInfo("Europe/Zurich")).strftime("%B %d %Y, %H:%M CET")

    # remove redundant "Executive Summary:" line from summary_text
    summary_text = re.sub(r"^(\*\*)?\s*Executive Summary\s*:?.*\n?", "", summary_text, flags=re.I)

    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        leftMargin=25 * mm,
        rightMargin=25 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()
    cover_title = ParagraphStyle("CoverTitle", parent=styles["Title"], fontName="Helvetica-Bold",
                                 fontSize=24, leading=28, spaceAfter=10, alignment=0)
    cover_label = ParagraphStyle("CoverLabel", parent=styles["BodyText"], fontName="Helvetica-Bold",
                                 fontSize=11, leading=14, spaceAfter=2, alignment=0)
    cover_text = ParagraphStyle("CoverText", parent=styles["BodyText"], fontName="Helvetica",
                                fontSize=11, leading=14, spaceAfter=8, alignment=0)
    heading = ParagraphStyle("Heading", parent=styles["Heading1"], fontName="Helvetica-Bold",
                             fontSize=14, leading=18, spaceAfter=10, spaceBefore=6)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica",
                          fontSize=11, leading=15, spaceAfter=10)
    list_item = ParagraphStyle("ListItem", parent=styles["BodyText"], fontName="Helvetica",
                               fontSize=11, leading=15, spaceAfter=10)

    def _md_to_html(text: str) -> str:
        t = html.escape(text)
        t = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", t)
        t = t.replace("\r\n", "\n").replace("\n", "<br/>")
        return t

    story = []
    if os.path.exists(LOGO_PATH):
        with PILImage.open(LOGO_PATH) as img:
            w, h = img.size
            target_width = 50 * mm
            target_height = target_width * (h / w)
        story.append(Image(LOGO_PATH, width=target_width, height=target_height, hAlign='LEFT'))
        story.append(Spacer(1, 10))

    story.append(Paragraph("Executive Summary", cover_title))
    story.append(HRFlowable(width="100%", thickness=0.6, color=colors.lightgrey))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Executive Summary Title:", cover_label))
    story.append(Paragraph(html.escape(exec_title), cover_text))
    story.append(Spacer(1, 6))
    story.append(Paragraph(f"Generated on: {html.escape(dt)}", cover_text))
    story.append(Paragraph(f"Source Document: {html.escape(src_name)}", cover_text))
    story.append(Paragraph("Prepared by: Transformate Consulting | AI Generated Summary", cover_text))

    # AI Processing Summary (if values provided)
    if all(v is not None for v in [prompt_tokens, completion_tokens, total_tokens, usd_cost, chf_cost]):
        story.append(Spacer(1, 8))
        story.append(Paragraph("AI Processing Summary:", cover_label))
        story.append(Paragraph(
            html.escape(f"Prompt Tokens: {prompt_tokens:,} | Completion Tokens: {completion_tokens:,} | Total: {total_tokens:,}"),
            cover_text
        ))
        story.append(Paragraph(
            html.escape(f"Estimated Cost: USD ${usd_cost:.4f} ≈ CHF {chf_cost:.4f}"),
            cover_text
        ))

    story.append(Spacer(1, 12))
    story.append(HRFlowable(width="100%", thickness=0.6, color=colors.lightgrey))
    story.append(Spacer(1, 8))
    story.append(Paragraph("Contact:", cover_label))
    story.append(Paragraph("Tony Ursich", cover_text))
    story.append(Paragraph("Chief Information Officer | Transformate Consulting", cover_text))
    story.append(Paragraph("E: tony.ursich@transformate.ch", cover_text))
    story.append(Paragraph("T: +41 76 577 1165", cover_text))
    story.append(Paragraph("W: www.transformate.ch", cover_text))
    story.append(PageBreak())

    # --- SUMMARY ---
    story.append(Paragraph(title, heading))
    for line in [l.strip() for l in summary_text.split("\n") if l.strip()]:
        clean_line = line.strip()
        if clean_line.startswith("**") and clean_line.endswith("**"):
            clean_line = clean_line.replace("**", "")
            story.append(Paragraph(f"<b>{html.escape(clean_line)}</b>", heading))
            story.append(Spacer(1, 6))
            continue
        if re.match(r"^[A-Z].+?:$", clean_line) or (
            "Purpose and Regulatory Intent" in clean_line
            or "Key Operational" in clean_line
            or "Recommended Next Steps" in clean_line
        ):
            story.append(Paragraph(f"<b>{html.escape(clean_line)}</b>", heading))
            story.append(Spacer(1, 8))
            continue
        if re.match(r"^\d+\.", clean_line):
            html_line = _md_to_html(clean_line)
            html_line = re.sub(r"</b>:", r":</b>", html_line)
            story.append(Paragraph(html_line, list_item))
            story.append(Spacer(1, 6))
            continue
        story.append(Paragraph(_md_to_html(clean_line), body))
        story.append(Spacer(1, 8))

    doc.build(story)


# ---------- FastAPI App ----------

app = FastAPI(title="Transformate Intelligence Suite (Dual Mode + Chatbase)", version="1.1")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # tighten for prod
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/")
async def health():
    return {"status": "ok", "service": "Transformate Intelligence Suite + Chatbase"}


@app.get("/files/{path:path}")
async def get_file(path: str):
    safe_path = os.path.normpath(path).lstrip(os.sep)
    full_path = os.path.join(OUTPUT_FOLDER, safe_path)
    if not os.path.isfile(full_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(full_path, media_type="application/pdf")


def _summarise_and_generate(pdf_path: str) -> dict:
    """Summarise a single PDF, generate a summary PDF, and push to Chatbase."""
    if not pdf_path.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail=f"Only PDF files are supported: {os.path.basename(pdf_path)}")

    text = extract_text_from_pdf(pdf_path)
    summary, pt, ct, tt, usd, chf = summarise_text(text, os.path.basename(pdf_path))

    # Output paths
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    output_pdf = os.path.join(OUTPUT_FOLDER, f"Summary_{base_name}.pdf")

    save_to_pdf(
        summary,
        output_pdf,
        source_document=pdf_path,
        prompt_tokens=pt,
        completion_tokens=ct,
        total_tokens=tt,
        usd_cost=usd,
        chf_cost=chf,
    )

    # === NEW: Push summary to Chatbase Knowledge API ===
    push_to_chatbase(base_name, summary)

    public_url = f"/files/{os.path.basename(output_pdf)}"

    return {
        "filename": os.path.basename(pdf_path),
        "summary_file": public_url,
        "token_usage": {"prompt": pt, "completion": ct, "total": tt},
        "cost": {"usd": usd, "chf": chf},
        "summary_preview": summary[:500] + ("..." if len(summary) > 500 else ""),
    }


@app.post("/summarise_demo")
async def summarise_demo():
    missing = [f for f in DEMO_FILES if not os.path.isfile(os.path.join(INPUT_FOLDER, f))]
    if missing:
        raise HTTPException(status_code=404, detail=f"Missing demo files: {missing}")

    results = []
    for name in DEMO_FILES:
        pdf_path = os.path.join(INPUT_FOLDER, name)
        results.append(_summarise_and_generate(pdf_path))

    return JSONResponse({"status": "success", "mode": "demo", "results": results})


@app.post("/summarise_upload")
async def summarise_upload(files: List[UploadFile] = File(...)):
    if not files:
        raise HTTPException(status_code=400, detail="Please upload at least one PDF.")
    if len(files) > 3:
        raise HTTPException(status_code=400, detail="Maximum 3 PDFs allowed per request.")

    results = []
    for uf in files:
        if not uf.filename.lower().endswith(".pdf"):
            raise HTTPException(status_code=400, detail=f"Only PDF files are supported: {uf.filename}")

        safe_name = re.sub(r"[^\w\-. ()]", "_", uf.filename)
        temp_path = os.path.join("/tmp", safe_name)
        with open(temp_path, "wb") as out:
            shutil.copyfileobj(uf.file, out)

        results.append(_summarise_and_generate(temp_path))

    return JSONResponse({"status": "success", "mode": "upload", "results": results})
